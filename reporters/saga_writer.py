from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

class SagaWriter:
    def __init__(self, scenario_name, output_dir):
        self.document = Document()
        
        # --- HEADER DOKUMEN ---
        section = self.document.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        title = self.document.add_heading(level=0)
        run = title.add_run(f'Heimdall Saga: {scenario_name}')
        run.font.name = 'Arial'
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(33, 33, 33) # Dark Grey
        
        self.output_dir = output_dir
        safe_name = "".join([c for c in scenario_name if c.isalnum() or c in (' ', '-', '_')]).strip()
        self.file_path = os.path.join(output_dir, f"Heimdall_Saga_{safe_name}.docx")

    def add_step(self, step_num, description, activity_name, screenshot_path, api_logs=None):
        """
        LAYOUT BARU: 1 STEP = 1 HALAMAN (Slide Style)
        Dijamin rapi dan tidak ada jarak aneh.
        """
        
        # 1. HEADER STEP
        # Tabel 1 baris untuk background header (opsional, biar rapi)
        table_head = self.document.add_table(rows=1, cols=1)
        table_head.autofit = False
        table_head.columns[0].width = Inches(6.5)
        cell = table_head.cell(0, 0)
        
        p = cell.paragraphs[0]
        run_step = p.add_run(f"STEP {step_num}")
        run_step.bold = True
        run_step.font.size = Pt(14)
        run_step.font.color.rgb = RGBColor(0, 51, 102) # Navy Blue
        
        # 2. ACTIVITY (Subtitle)
        p_act = self.document.add_paragraph()
        p_act.paragraph_format.space_after = Pt(6)
        run_act = p_act.add_run(f"📍 Position: {activity_name}")
        run_act.font.size = Pt(9)
        run_act.italic = True
        run_act.font.color.rgb = RGBColor(128, 128, 128)

        # 3. NARASI STORYTELLER (PENTING)
        # Font Serif (Georgia) agar kontras dengan header
        p_desc = self.document.add_paragraph(description)
        p_desc.paragraph_format.space_after = Pt(12)
        p_desc.style.font.name = 'Georgia'
        p_desc.style.font.size = Pt(12)

        # 4. SCREENSHOT (BESAR & TENGAH)
        if screenshot_path and os.path.exists(screenshot_path):
            try:
                p_img = self.document.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                # Lebar 3.0 Inch (Cukup besar tapi muat di halaman)
                run_img.add_picture(screenshot_path, width=Inches(3.0))
            except:
                pass

        # 5. TABEL API (JIKA ADA)
        if api_logs and len(api_logs) > 0:
            p_api = self.document.add_paragraph()
            p_api.paragraph_format.space_before = Pt(12)
            run_api = p_api.add_run("📡 Network Activity Detected:")
            run_api.bold = True
            run_api.font.size = Pt(9)
            run_api.font.name = 'Arial'

            # Tabel Grid
            table = self.document.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            table.autofit = False
            
            # Lebar Kolom
            table.columns[0].width = Inches(0.7) # Method
            table.columns[1].width = Inches(4.5) # Endpoint (Lebar)
            table.columns[2].width = Inches(1.2) # Status

            # Header
            hdr = table.rows[0].cells
            self._style_cell(hdr[0], "METHOD", bg="E0E0E0", bold=True)
            self._style_cell(hdr[1], "ENDPOINT RESOURCE", bg="E0E0E0", bold=True)
            self._style_cell(hdr[2], "STATUS", bg="E0E0E0", bold=True)

            for log in api_logs:
                method = log.get('method', '-')
                endpoint = log.get('endpoint', '-')
                status = str(log.get('status', '-'))
                
                # Potong endpoint super panjang
                if len(endpoint) > 60: endpoint = "..." + endpoint[-55:]

                row = table.add_row().cells
                self._style_cell(row[0], method, size=8)
                self._style_cell(row[1], endpoint, size=8, font='Consolas')
                
                icon = "✅" if status.startswith('2') else "❌"
                if status == '-': icon = "⏳" # Loading icon
                
                self._style_cell(row[2], f"{icon} {status}", size=8, bold=True)

        # 6. PAGE BREAK (KUNCI RAPIH)
        # Ganti halaman setiap selesai satu langkah.
        self.document.add_page_break()

    def _style_cell(self, cell, text, bold=False, size=9, font='Arial', bg=None):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.bold = bold
        run.font.size = Pt(size)
        run.font.name = font
        
        if bg:
            tc_pr = cell._element.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg)
            tc_pr.append(shd)

    def save(self):
        try:
            self.document.save(self.file_path)
            print(f"  📄 Report saved: {self.file_path}")
        except Exception as e:
            print(f"!!! Error saving docx: {e}")